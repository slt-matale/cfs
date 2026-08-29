import os
import secrets

from fastapi import FastAPI, HTTPException, Header, Depends
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel

from backend.database import supabase
from backend.chatbot import questions
from backend.gemini_client import generate_improvement, generate_negative_summary

from collections import Counter
from datetime import datetime, timedelta
from pathlib import Path
from tempfile import gettempdir

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH



app = FastAPI(
    title="AI Customer Feedback System",
    version="2.0"
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.get("/")
def home():
    return {
        "message": "AI Customer Feedback Bot Running"
    }


@app.get("/questions")
def get_questions():
    return questions



class LoginRequest(BaseModel):
    username: str
    password: str


ADMIN_USERNAME = os.environ.get("ADMIN_USERNAME", "admin")
ADMIN_PASSWORD = os.environ.get("ADMIN_PASSWORD", "admin123")

# A fresh random token is generated each time the server starts. The
# frontend receives it on successful login and must send it back as
# "Authorization: Bearer <token>" on every protected request below.
# This means the hardcoded admin/admin123 check in the old client-side
# JS can no longer grant access on its own - the backend is now the
# only thing that decides who's logged in.
ADMIN_TOKEN = secrets.token_urlsafe(32)


def require_admin(authorization: str | None = Header(default=None)):
    """Dependency that protects admin-only endpoints with a bearer token."""

    if not authorization or not authorization.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Not authenticated")

    token = authorization.removeprefix("Bearer ").strip()

    if not secrets.compare_digest(token, ADMIN_TOKEN):
        raise HTTPException(status_code=401, detail="Invalid or expired session")


@app.post("/login")
def login(data: LoginRequest):

   
    if data.username == ADMIN_USERNAME and data.password == ADMIN_PASSWORD:

        return {
            "success": True,
            "message": "Login successful",
            "token": ADMIN_TOKEN
        }

    raise HTTPException(
        status_code=401,
        detail="Invalid username or password"
    )


class Feedback(BaseModel):

    service: str | None = None
    phone: str | None = None
    waiting: str | None = None
    staff: str | None = None
    office: str | None = None
    parking: str | None = None
    comment: str | None = None



def calculate_sentiment(feedback):

    ratings = [
        feedback.waiting,
        feedback.staff,
        feedback.office,
        feedback.parking
    ]

    positive = 0
    neutral = 0
    negative = 0

    for rating in ratings:

        if rating in ["Excellent", "Good"]:
            positive += 1

        elif rating == "Average":
            neutral += 1

        elif rating in ["Poor", "Very Poor"]:
            negative += 1

    # Comment can also influence the result
    comment = (feedback.comment or "").strip().lower()

    positive_words = [
        "good",
        "great",
        "excellent",
        "friendly",
        "helpful",
        "satisfied",
        "happy",
        "nice",
        "perfect",
        "thank"
    ]

    negative_words = [
        "bad",
        "poor",
        "very poor",
        "terrible",
        "worst",
        "problem",
        "issue",
        "improve",
        "improvement",
        "slow",
        "unhappy"
    ]

    if comment:

        if any(word in comment for word in negative_words):
            negative += 2

        elif any(word in comment for word in positive_words):
            positive += 2

    if positive > negative and positive > neutral:
        return "Positive"

    if negative > positive and negative > neutral:
        return "Negative"

    return "Neutral"


@app.post("/feedback")
def save_feedback(feedback: Feedback):

    try:

        sentiment = calculate_sentiment(feedback)

        data = {
            "service": feedback.service,
            "phone": feedback.phone,
            "waiting": feedback.waiting,
            "staff": feedback.staff,
            "office": feedback.office,
            "parking": feedback.parking,
            "comment": feedback.comment,
            "sentiment": sentiment
        }

        response = (
            supabase
            .table("customer_feedback")
            .insert(data)
            .execute()
        )

        return {
            "success": True,
            "message": "Feedback saved successfully",
            "sentiment": sentiment,
            "data": response.data
        }

    except Exception as e:

        print("SUPABASE SAVE ERROR:", e)

        raise HTTPException(
            status_code=500,
            detail=str(e)
        )

@app.get("/test-supabase")
def test_supabase():

    try:

        response = (
            supabase
            .table("customer_feedback")
            .select("*")
            .limit(1)
            .execute()
        )

        return {
            "success": True,
            "message": "Supabase connection is working",
            "data": response.data
        }

    except Exception as e:

        raise HTTPException(
            status_code=500,
            detail=str(e)
        )

def get_all_feedback(
    start_date: str | None = None,
    end_date: str | None = None,
    sentiment: str | None = None
):

    query = (
        supabase
        .table("customer_feedback")
        .select("*")
    )

    # created_at is stored as a timestamp, so "end_date" needs to include
    # the whole day, not stop at midnight - otherwise feedback submitted
    # later on the end date would be silently excluded.
    if start_date:
        query = query.gte("created_at", f"{start_date}T00:00:00")

    if end_date:
        query = query.lte("created_at", f"{end_date}T23:59:59")

    if sentiment:
        query = query.eq("sentiment", sentiment)

    response = query.execute()

    return response.data or []


@app.get("/feedback-date-range")
def feedback_date_range(_: None = Depends(require_admin)):
    """Returns the earliest and latest feedback dates available, so the
    frontend can bound its date-range picker to real data instead of
    letting the admin pick a range that has nothing in it."""

    try:

        earliest = (
            supabase
            .table("customer_feedback")
            .select("created_at")
            .order("created_at", desc=False)
            .limit(1)
            .execute()
        )

        earliest_date = None

        if earliest.data:
            earliest_date = str(earliest.data[0]["created_at"])[:10]

        today = datetime.now().date().isoformat()

        return {
            "earliest_date": earliest_date or today,
            "latest_date": today
        }

    except Exception as e:

        print("DATE RANGE ERROR:", e)

        raise HTTPException(
            status_code=500,
            detail=str(e)
        )


@app.get("/negative-feedback")
def negative_feedback(
    start_date: str | None = None,
    end_date: str | None = None,
    _: None = Depends(require_admin)
):
    """All negative-sentiment feedback in the given range, newest first."""

    try:

        feedback = get_all_feedback(
            start_date,
            end_date,
            sentiment="Negative"
        )

        feedback.sort(
            key=lambda item: item.get("created_at") or "",
            reverse=True
        )

        return {
            "count": len(feedback),
            "feedback": feedback
        }

    except Exception as e:

        print("NEGATIVE FEEDBACK ERROR:", e)

        raise HTTPException(
            status_code=500,
            detail=str(e)
        )


class ImproveRequest(BaseModel):

    service: str | None = None
    phone: str | None = None
    waiting: str | None = None
    staff: str | None = None
    office: str | None = None
    parking: str | None = None
    comment: str | None = None


@app.post("/negative-feedback/improve")
def improve_feedback(
    data: ImproveRequest,
    _: None = Depends(require_admin)
):
    """AI-generated improvement suggestion for one specific negative
    feedback entry - called on demand as the admin browses the list,
    not pre-generated in bulk."""

    try:

        suggestion = generate_improvement(data.model_dump())

        return {"suggestion": suggestion}

    except Exception as e:

        print("GEMINI IMPROVE ERROR:", e)

        raise HTTPException(
            status_code=502,
            detail=str(e)
        )


@app.get("/negative-feedback/summary")
def negative_feedback_summary(
    start_date: str | None = None,
    end_date: str | None = None,
    _: None = Depends(require_admin)
):
    """One AI-generated overview of the main complaint themes and top
    recommended improvements across all negative feedback in range."""

    try:

        feedback = get_all_feedback(
            start_date,
            end_date,
            sentiment="Negative"
        )

        if not feedback:

            return {
                "summary": (
                    "No negative feedback found for the selected period."
                ),
                "count": 0
            }

        summary = generate_negative_summary(feedback)

        return {
            "summary": summary,
            "count": len(feedback)
        }

    except HTTPException:

        raise

    except Exception as e:

        print("GEMINI SUMMARY ERROR:", e)

        raise HTTPException(
            status_code=502,
            detail=str(e)
        )


@app.get("/dashboard-data")
def dashboard_data(
    start_date: str | None = None,
    end_date: str | None = None,
    _: None = Depends(require_admin)
):

    try:

        feedback = get_all_feedback(start_date, end_date)

        total = len(feedback)

        positive = 0
        neutral = 0
        negative = 0

        service = Counter()
        waiting = Counter()
        office = Counter()
        staff = Counter()
        parking = Counter()

        today = datetime.now().date()

        start_of_week = (
            today -
            timedelta(days=today.weekday())
        )

        weekly = {
            "Monday": 0,
            "Tuesday": 0,
            "Wednesday": 0,
            "Thursday": 0,
            "Friday": 0,
            "Saturday": 0,
            "Sunday": 0
        }

        for item in feedback:

            sentiment = (
                item.get("sentiment")
                or "Neutral"
            )

            if sentiment == "Positive":
                positive += 1

            elif sentiment == "Negative":
                negative += 1

            else:
                neutral += 1

            if item.get("service"):
                service[item["service"]] += 1

            if item.get("waiting"):
                waiting[item["waiting"]] += 1

            if item.get("office"):
                office[item["office"]] += 1

            if item.get("staff"):
                staff[item["staff"]] += 1

            if item.get("parking"):
                parking[item["parking"]] += 1

    
            created_at = item.get("created_at")

            if created_at:

                try:

                    date_text = str(created_at)

                    if date_text.endswith("Z"):
                        date_text = date_text[:-1]

                    created_date = (
                        datetime
                        .fromisoformat(date_text)
                        .date()
                    )

                    if (
                        start_of_week
                        <= created_date
                        <= today
                    ):

                        day_name = (
                            created_date
                            .strftime("%A")
                        )

                        weekly[day_name] += 1

                except Exception:
                    pass

        return {
            "total_feedback": total,

            "sentiment": {
                "positive": positive,
                "neutral": neutral,
                "negative": negative
            },

            "service": dict(service),

            "waiting": dict(waiting),

            "office": dict(office),

            "staff": dict(staff),

            "parking": dict(parking),

            "weekly": weekly,

            "feedback": feedback
        }

    except Exception as e:

        print("DASHBOARD ERROR:", e)

        raise HTTPException(
            status_code=500,
            detail=str(e)
        )

@app.get("/ai-summary")
def ai_summary(
    start_date: str | None = None,
    end_date: str | None = None,
    _: None = Depends(require_admin)
):

    try:

        feedback = get_all_feedback(start_date, end_date)

        total = len(feedback)

        positive = sum(
            1
            for item in feedback
            if item.get("sentiment") == "Positive"
        )

        negative = sum(
            1
            for item in feedback
            if item.get("sentiment") == "Negative"
        )

        neutral = (
            total -
            positive -
            negative
        )

        def percentage(value):

            if total == 0:
                return 0

            return round(
                (value / total) * 100,
                1
            )

        positive_percentage = percentage(positive)
        neutral_percentage = percentage(neutral)
        negative_percentage = percentage(negative)

        service = Counter(
            item.get("service")
            for item in feedback
            if item.get("service")
        )

        waiting = Counter(
            item.get("waiting")
            for item in feedback
            if item.get("waiting")
        )

        staff = Counter(
            item.get("staff")
            for item in feedback
            if item.get("staff")
        )

        office = Counter(
            item.get("office")
            for item in feedback
            if item.get("office")
        )
        comments = []

        for item in feedback:

            comment = str(
                item.get("comment") or ""
            ).strip()

            if comment:

                comments.append({
                    "comment": comment,
                    "sentiment":
                        item.get("sentiment")
                        or "Neutral"
                })

        if total == 0:

            overall_summary = (
                "There is currently no customer feedback "
                "available for analysis."
            )

        elif positive_percentage >= 60:

            overall_summary = (
                "Overall customer feedback is positive. "
                "Most customers reported satisfactory "
                "experiences. The positive feedback indicates "
                "that the current customer experience is "
                "generally effective, while negative and "
                "neutral responses identify areas that can "
                "still be improved."
            )

        elif positive_percentage >= 40:

            overall_summary = (
                "Overall customer feedback is mixed but "
                "generally satisfactory. Positive experiences "
                "are present, while neutral and negative "
                "feedback indicates areas where customer "
                "service can be improved."
            )

        else:

            overall_summary = (
                "Overall customer feedback indicates that "
                "improvement is required. Negative and neutral "
                "feedback should be reviewed to identify the "
                "main causes of customer dissatisfaction."
            )

        key_findings = []

        if total > 0:

            key_findings.append(
                f"{positive} customers provided positive "
                f"feedback ({positive_percentage}%)."
            )

            key_findings.append(
                f"{negative} customers provided negative "
                f"feedback ({negative_percentage}%) "
                f"and should be reviewed."
            )

            key_findings.append(
                f"{neutral} customers provided neutral "
                f"feedback ({neutral_percentage}%)."
            )

        # Lowest rating categories
        def lowest_rating(counter):

            if not counter:
                return None

            negative_values = {
                "Very Poor": 5,
                "Poor": 4,
                "Average": 3,
                "Good": 2,
                "Excellent": 1
            }

            available = [
                (rating, count)
                for rating, count in counter.items()
                if rating in negative_values
            ]

            if not available:
                return None

            # Prioritize Very Poor/Poor
            available.sort(
                key=lambda x: (
                    negative_values.get(x[0], 99),
                    -x[1]
                ),
                reverse=False
            )

            for rating, count in available:

                if rating in ["Very Poor", "Poor"]:
                    return rating, count

            return None

        staff_issue = lowest_rating(staff)

        if staff_issue:

            key_findings.append(
                f"Staff rating area requiring attention: "
                f"{staff_issue[0]} ({staff_issue[1]})."
            )

        waiting_issue = lowest_rating(waiting)

        if waiting_issue:

            key_findings.append(
                f"Waiting-time rating requiring attention: "
                f"{waiting_issue[0]} ({waiting_issue[1]})."
            )

        office_issue = lowest_rating(office)

        if office_issue:

            key_findings.append(
                f"Office-environment rating requiring attention: "
                f"{office_issue[0]} ({office_issue[1]})."
            )

        recommendations = []

        if negative > 0:

            recommendations.append(
                "Review negative customer comments and "
                "identify recurring service problems."
            )

        if waiting.get("Poor", 0) > 0 or waiting.get("Very Poor", 0) > 0:

            recommendations.append(
                "Improve waiting-time management and "
                "reduce customer waiting periods."
            )

        if staff.get("Poor", 0) > 0 or staff.get("Very Poor", 0) > 0:

            recommendations.append(
                "Provide additional staff training and "
                "customer-service support."
            )

        if office.get("Poor", 0) > 0 or office.get("Very Poor", 0) > 0:

            recommendations.append(
                "Review the office environment and identify "
                "areas requiring improvement."
            )

        if not recommendations:

            recommendations.append(
                "Continue monitoring customer feedback "
                "regularly to maintain service quality."
            )

        return {

            "title":
                "AI Customer Feedback Summary",

            "company":
                "SLTMobitel Customer Experience",

            "generated_at":
                datetime.now().strftime(
                    "%Y-%m-%d %H:%M:%S"
                ),

            "total_feedback":
                total,

            "positive":
                positive,

            "neutral":
                neutral,

            "negative":
                negative,

            "positive_percentage":
                positive_percentage,

            "neutral_percentage":
                neutral_percentage,

            "negative_percentage":
                negative_percentage,

            "service":
                dict(service),

            "waiting":
                dict(waiting),

            "staff":
                dict(staff),

            "office":
                dict(office),

            "overall_summary":
                overall_summary,

            "summary":
                overall_summary,

            "key_findings":
                key_findings,

            "recommendations":
                recommendations,

            "comments":
                comments
        }

    except Exception as e:

        print("AI SUMMARY ERROR:", e)

        raise HTTPException(
            status_code=500,
            detail=str(e)
        )

def create_title(document, text):

    paragraph = document.add_paragraph()

    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run(text)

    run.bold = True
    run.font.size = Pt(20)


def add_bullet(document, text):

    document.add_paragraph(
        text,
        style="List Bullet"
    )

@app.get("/download-ai-summary")
def download_ai_summary(
    start_date: str | None = None,
    end_date: str | None = None,
    _: None = Depends(require_admin)
):

    try:

        data = ai_summary(start_date, end_date, _)

        document = Document()

        create_title(
            document,
            "AI Customer Feedback Summary"
        )

        subtitle = document.add_paragraph()

        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = subtitle.add_run(
            "SLTMobitel Customer Experience"
        )

        run.bold = True

        generated = document.add_paragraph()

        generated.alignment = WD_ALIGN_PARAGRAPH.CENTER

        generated.add_run(
            "Report Generated: "
            + data["generated_at"]
        )

        document.add_heading(
            "1. Feedback Overview",
            level=1
        )

        document.add_paragraph(
            f"Total Feedback: "
            f"{data['total_feedback']}"
        )

        document.add_paragraph(
            f"Positive Feedback: "
            f"{data['positive']}"
        )

        document.add_paragraph(
            f"Neutral Feedback: "
            f"{data['neutral']}"
        )

        document.add_paragraph(
            f"Negative Feedback: "
            f"{data['negative']}"
        )


        document.add_heading(
            "2. Sentiment Analysis",
            level=1
        )

        document.add_paragraph(
            f"Positive: "
            f"{data['positive_percentage']}%"
        )

        document.add_paragraph(
            f"Neutral: "
            f"{data['neutral_percentage']}%"
        )

        document.add_paragraph(
            f"Negative: "
            f"{data['negative_percentage']}%"
        )

        document.add_heading(
            "3. Service Summary",
            level=1
        )

        for name, count in data["service"].items():

            document.add_paragraph(
                f"{name}: {count}"
            )


        document.add_heading(
            "4. Customer Experience Ratings",
            level=1
        )

        document.add_paragraph(
            "Waiting Time:"
        )

        for name, count in data["waiting"].items():

            document.add_paragraph(
                f"{name}: {count}"
            )

        document.add_paragraph(
            "Staff:"
        )

        for name, count in data["staff"].items():

            document.add_paragraph(
                f"{name}: {count}"
            )

        document.add_paragraph(
            "Office Environment:"
        )

        for name, count in data["office"].items():

            document.add_paragraph(
                f"{name}: {count}"
            )


        document.add_heading(
            "5. Overall Summary",
            level=1
        )

        document.add_paragraph(
            data["overall_summary"]
        )


        document.add_heading(
            "6. Key Findings",
            level=1
        )

        for finding in data["key_findings"]:

            add_bullet(
                document,
                finding
            )


        document.add_heading(
            "7. Recommendations",
            level=1
        )

        for recommendation in data["recommendations"]:

            add_bullet(
                document,
                recommendation
            )


        document.add_heading(
            "8. Customer Comments",
            level=1
        )

        if data["comments"]:

            for item in data["comments"]:

                paragraph = document.add_paragraph()

                paragraph.add_run(
                    "Comment: "
                ).bold = True

                paragraph.add_run(
                    item["comment"]
                )

                paragraph.add_run(
                    "\nSentiment: "
                    + item["sentiment"]
                )

        else:

            document.add_paragraph(
                "No customer comments available."
            )

        filename = (
            "AI_Customer_Feedback_Summary.docx"
        )

        path = (
            Path(gettempdir())
            / filename
        )

        document.save(path)

        return FileResponse(
            path=str(path),
            filename=filename,
            media_type=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            )
        )

    except Exception as e:

        print("AI SUMMARY WORD ERROR:", e)

        raise HTTPException(
            status_code=500,
            detail=str(e)
        )


@app.get("/download-customer-feedback")
def download_customer_feedback(
    start_date: str | None = None,
    end_date: str | None = None,
    _: None = Depends(require_admin)
):

    try:

        feedback = get_all_feedback(start_date, end_date)

        document = Document()

        create_title(
            document,
            "Customer Feedback Data"
        )

        subtitle = document.add_paragraph()

        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle.add_run(
            "SLTMobitel Customer Experience"
        ).bold = True

        document.add_paragraph(
            f"Total Records: {len(feedback)}"
        )

        document.add_paragraph(
            "Generated: "
            + datetime.now().strftime(
                "%Y-%m-%d %H:%M:%S"
            )
        )


        document.add_heading(
            "Customer Feedback Records",
            level=1
        )

        headers = [
            "ID",
            "Service",
            "Phone",
            "Waiting",
            "Staff",
            "Office",
            "Parking",
            "Comment",
            "Sentiment",
            "Created At"
        ]

        table = document.add_table(
            rows=1,
            cols=len(headers)
        )

        table.style = "Table Grid"

        for i, header in enumerate(headers):

            table.rows[0].cells[i].text = header

        for item in feedback:

            row = table.add_row().cells

            values = [

                item.get("id", ""),

                item.get("service", ""),

                item.get("phone", ""),

                item.get("waiting", ""),

                item.get("staff", ""),

                item.get("office", ""),

                item.get("parking", ""),

                item.get("comment", ""),

                item.get("sentiment", ""),

                item.get("created_at", "")
            ]

            for i, value in enumerate(values):

                row[i].text = str(
                    value or ""
                )

        filename = (
            "Customer_Feedback_Data.docx"
        )

        path = (
            Path(gettempdir())
            / filename
        )

        document.save(path)

        return FileResponse(
            path=str(path),
            filename=filename,
            media_type=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            )
        )

    except Exception as e:

        print(
            "CUSTOMER FEEDBACK WORD ERROR:",
            e
        )

        raise HTTPException(
            status_code=500,
            detail=str(e)
        )